"""Exchange-certificate parser (PDF/PNG/DOCX).

Extracts deal_id, isin, amount, date via regex.
Uses Tesseract OCR for images and scanned PDFs.
"""
from __future__ import annotations
from pathlib import Path
from typing import Any, Dict, List, Optional
import re
from loguru import logger

DEAL_RE = re.compile(r"(?:сделка|deal)[\s#№]*(\d+)", re.IGNORECASE)
AMOUNT_RE = re.compile(r"(\d[\d\s]*(?:\.\d{2})?)\s*(?:тенге|KZT|USD)", re.IGNORECASE)
ISIN_RE = re.compile(r"([A-Z]{2}[A-Z0-9]{9}\d)")
DATE_RE = re.compile(r"(\d{2}[./]\d{2}[./]\d{4})")


def parse_certificate(file_path: Path | str) -> Dict[str, Any]:
    p = Path(file_path)
    text = _extract_text(p)
    if not text:
        return {"deals": [], "warnings": ["No text extracted"]}

    deals = []
    for deal_id in DEAL_RE.findall(text):
        deals.append({"deal_id": deal_id, "source": "regex"})
    amounts = AMOUNT_RE.findall(text)
    isins = ISIN_RE.findall(text)
    dates = DATE_RE.findall(text)
    return {
        "filename": p.name,
        "deals": deals,
        "amounts_found": amounts,
        "isins_found": isins,
        "dates_found": dates,
        "text_sample": text[:500],
    }


def _extract_text(p: Path) -> Optional[str]:
    suffix = p.suffix.lower()
    if suffix == ".docx":
        return _extract_docx_text(p)
    if suffix == ".pdf":
        return _extract_pdf_text(p)
    if suffix in (".png", ".jpg", ".jpeg"):
        return _ocr_image(p)
    return None


def _extract_docx_text(p: Path) -> Optional[str]:
    try:
        from docx import Document
        doc = Document(p)
        parts: List[str] = []
        # Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        # Tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts) if parts else None
    except Exception:
        logger.warning(f"docx read failed: {p}")
        return None


def _extract_pdf_text(p: Path) -> Optional[str]:
    try:
        import pdfplumber
        parts: List[str] = []
        with pdfplumber.open(p) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ""
                if text.strip():
                    parts.append(text)
                else:
                    # Scanned PDF — fallback to OCR
                    logger.info(f"Page text empty, trying OCR for: {p}")
                    try:
                        img_text = _ocr_pdf_page(page)
                        if img_text:
                            parts.append(img_text)
                    except Exception as e:
                        logger.debug(f"OCR failed for page: {e}")
        return "\n".join(parts) if parts else None
    except Exception:
        logger.warning(f"pdf read failed: {p}")
        return None


def _configure_tesseract():
    import shutil, os
    tess = shutil.which("tesseract")
    if tess:
        return
    for p in [
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Tesseract-OCR\tesseract.exe",
    ]:
        if os.path.exists(p):
            import pytesseract
            pytesseract.pytesseract.tesseract_cmd = p
            break


def _ocr_image(p: Path) -> Optional[str]:
    try:
        import pytesseract
        from PIL import Image
        _configure_tesseract()
        img = Image.open(p)
        text = pytesseract.image_to_string(img, lang="rus+kaz+eng")
        return text if text.strip() else None
    except Exception as e:
        logger.warning(f"OCR image failed: {p} — {e}")
        return None


def _ocr_pdf_page(page) -> Optional[str]:
    """Convert PDF page to image and OCR."""
    try:
        import pytesseract
        from PIL import Image
        _configure_tesseract()
        img = page.to_image(resolution=300).original
        text = pytesseract.image_to_string(img, lang="rus+kaz+eng")
        return text if text.strip() else None
    except Exception as e:
        logger.warning(f"OCR PDF page failed: {e}")
        return None
